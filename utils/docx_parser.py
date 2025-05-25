#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Word文档解析工具 - 处理Word文档内容提取
"""

import os
import re

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


class DocxParser:
    """Word文档解析器"""

    def __init__(self, file_path):
        self.file_path = file_path
        self.document = None
        self.paragraphs = []
        self.tables = []

        self._load_document()

    def _load_document(self):
        """加载文档"""
        try:
            if not os.path.exists(self.file_path):
                raise FileNotFoundError(f"文件不存在: {self.file_path}")

            self.document = Document(self.file_path)

            # 提取段落
            self.paragraphs = [p.text for p in self.document.paragraphs]

            # 提取表格
            self.tables = []
            for table in self.document.tables:
                table_data = []
                for i, row in enumerate(table.rows):
                    row_data = []
                    for j, cell in enumerate(row.cells):
                        row_data.append(cell.text)
                    table_data.append(row_data)
                self.tables.append(table_data)

        except PackageNotFoundError:
            raise ValueError(f"无法打开文件，可能不是有效的Word文档: {self.file_path}")
        except Exception as e:
            raise ValueError(f"加载文档时出错: {str(e)}")

    def extract_with_regex(self, pattern, group=0):
        """使用正则表达式提取文本"""
        try:
            # 将所有段落合并为一个文本
            all_text = "\n".join(self.paragraphs)

            # 应用正则表达式
            regex = re.compile(pattern)
            matches = regex.findall(all_text)

            # 处理结果
            if not matches:
                return ""

            if isinstance(matches[0], tuple) and len(matches[0]) > group:
                # 如果匹配结果是元组（有多个捕获组）
                result = [m[group] for m in matches]
                return "\n".join(result)
            else:
                # 如果匹配结果是字符串（单个捕获组或整个匹配）
                return "\n".join(matches)

        except Exception as e:
            return f"正则表达式错误: {str(e)}"

    def extract_by_position(self, start_index, end_index=None):
        """通过位置索引提取文本"""
        try:
            if start_index < 0 or start_index >= len(self.paragraphs):
                return ""

            if end_index is None or end_index < 0:
                end_index = len(self.paragraphs)

            end_index = min(end_index, len(self.paragraphs))

            return "\n".join(self.paragraphs[start_index:end_index])

        except Exception as e:
            return f"位置提取错误: {str(e)}"

    def extract_by_bookmark(self, bookmark_name):
        """通过书签提取文本"""
        try:
            # 遍历文档中的所有书签
            if not hasattr(self.document, 'element') or not hasattr(self.document.element, 'xpath'):
                return "文档不支持书签访问"

            bookmarks = {}
            for bookmark_elem in self.document.element.xpath('//w:bookmark'):
                name = bookmark_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')
                if name == bookmark_name:
                    # 找到书签，提取其内容（这部分需要根据具体情况实现，这里是简化版本）
                    # 实际应用可能需要更复杂的解析逻辑
                    return f"书签 '{bookmark_name}' 的内容"

            return f"未找到书签: {bookmark_name}"

        except Exception as e:
            return f"书签提取错误: {str(e)}"

    def extract_table_cell(self, table_index, row_index, col_index):
        """提取表格单元格内容"""
        try:
            if table_index < 0 or table_index >= len(self.tables):
                return f"表格索引越界: {table_index}"

            table = self.tables[table_index]

            if row_index < 0 or row_index >= len(table):
                return f"行索引越界: {row_index}"

            if col_index < 0 or col_index >= len(table[row_index]):
                return f"列索引越界: {col_index}"

            return table[row_index][col_index]

        except Exception as e:
            return f"表格单元格提取错误: {str(e)}"

    def extract_table_column(self, table_index, col_index, has_header=True):
        """提取表格列"""
        try:
            if table_index < 0 or table_index >= len(self.tables):
                return f"表格索引越界: {table_index}"

            table = self.tables[table_index]

            if len(table) == 0:
                return "表格为空"

            if col_index < 0 or col_index >= len(table[0]):
                return f"列索引越界: {col_index}"

            # 提取列数据
            start_row = 1 if has_header else 0
            column_data = [row[col_index] for row in table[start_row:]]

            return "\n".join(column_data)

        except Exception as e:
            return f"表格列提取错误: {str(e)}"

    def extract_table_row(self, table_index, row_index):
        """提取表格行"""
        try:
            if table_index < 0 or table_index >= len(self.tables):
                return f"表格索引越界: {table_index}"

            table = self.tables[table_index]

            if row_index < 0 or row_index >= len(table):
                return f"行索引越界: {row_index}"

            # 提取行数据
            row_data = table[row_index]

            return "\t".join(row_data)

        except Exception as e:
            return f"表格行提取错误: {str(e)}"

    def extract_table(self, table_index, has_header=True):
        """提取整个表格"""
        try:
            if table_index < 0 or table_index >= len(self.tables):
                return f"表格索引越界: {table_index}"

            table = self.tables[table_index]

            if not table:
                return "表格为空"

            # 返回表格数据，适合测试显示
            if has_header and len(table) > 0:
                # 如果有表头，则第一行作为表头显示
                return table
            else:
                # 无表头
                return table

        except Exception as e:
            return f"表格提取错误: {str(e)}"
