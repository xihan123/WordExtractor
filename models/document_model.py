#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文档内容模型 - 处理Word文档内容
"""

import os
import re

from PyQt6.QtCore import QObject, pyqtSignal, pyqtSlot, QRunnable, QThreadPool
from docx import Document


class DocumentContent:
    """表示文档内容的类"""

    def __init__(self, file_path=None):
        self.file_path = file_path
        self.title = ""
        self.author = ""
        self.created_date = ""
        self.last_modified_date = ""

        # 文档内容结构
        self.paragraphs = []  # 段落列表
        self.tables = []  # 表格列表

        # 原始文档对象
        self.document = None

        # 加载状态
        self.is_loaded = False
        self.load_error = ""

        if file_path and os.path.exists(file_path):
            try:
                self._load_document()
            except Exception as e:
                self.load_error = str(e)

    def _load_document(self):
        """加载文档内容"""
        try:
            self.document = Document(self.file_path)

            # 提取文档元数据
            core_props = self.document.core_properties
            self.title = core_props.title or os.path.basename(self.file_path)
            self.author = core_props.author or "未知"
            self.created_date = core_props.created or "未知"
            self.last_modified_date = core_props.modified or "未知"

            # 提取段落
            self.paragraphs = [
                {
                    "index": i,
                    "text": p.text,
                    "style": p.style.name,
                    "level": p.paragraph_format.alignment
                }
                for i, p in enumerate(self.document.paragraphs) if p.text.strip()
            ]

            # 提取表格
            for i, table in enumerate(self.document.tables):
                table_data = []
                for r_idx, row in enumerate(table.rows):
                    row_data = []
                    for c_idx, cell in enumerate(row.cells):
                        row_data.append({
                            "text": cell.text,
                            "position": (r_idx, c_idx)
                        })
                    table_data.append(row_data)

                self.tables.append({
                    "index": i,
                    "rows": len(table.rows),
                    "cols": len(table.columns),
                    "data": table_data
                })

            self.is_loaded = True
            return True

        except Exception as e:
            self.load_error = f"加载文档出错: {str(e)}"
            return False

    def get_table_as_html(self, table_index):
        """将表格转换为HTML格式以便显示"""
        if not self.is_loaded or table_index >= len(self.tables):
            return "<p>表格不可用</p>"

        table = self.tables[table_index]

        html = "<table border='1' cellpadding='3' style='border-collapse: collapse;'>"

        for row in table["data"]:
            html += "<tr>"
            for cell in row:
                html += f"<td>{cell['text']}</td>"
            html += "</tr>"

        html += "</table>"
        return html

    def get_document_html(self):
        """将文档内容转换为HTML格式以便显示"""
        if not self.is_loaded:
            return f"<p>文档加载失败: {self.load_error}</p>"

        html = f"<h1>{self.title}</h1>"
        html += f"<p><b>作者:</b> {self.author}</p>"
        html += f"<p><b>创建日期:</b> {self.created_date}</p>"
        html += f"<p><b>最后修改:</b> {self.last_modified_date}</p>"
        html += "<hr/>"

        # 添加段落
        for para in self.paragraphs:
            style = ""

            # 根据段落样式应用不同的HTML样式
            if para["style"].startswith("Heading"):
                level = int(para["style"].replace("Heading ", "")) if para["style"] != "Heading" else 1
                if 1 <= level <= 6:
                    html += f"<h{level}>{para['text']}</h{level}>"
                    continue

            html += f"<p>{para['text']}</p>"

        # 添加表格
        for i, table in enumerate(self.tables):
            html += f"<h3>表格 {i + 1}</h3>"
            html += self.get_table_as_html(i)

        return html

    def extract_text_with_regex(self, pattern, group=0):
        """使用正则表达式从文档中提取文本"""
        if not self.is_loaded:
            return None

        all_text = "\n".join([p["text"] for p in self.paragraphs])

        try:
            regex = re.compile(pattern)
            matches = regex.findall(all_text)

            if matches:
                if isinstance(matches[0], tuple) and len(matches[0]) > group:
                    return [m[group] for m in matches]
                return matches
            return []
        except Exception as e:
            return [f"正则表达式错误: {str(e)}"]

    def extract_text_by_index(self, start_idx, end_idx=None):
        """通过索引位置从文档中提取文本"""
        if not self.is_loaded:
            return None

        if start_idx >= len(self.paragraphs):
            return None

        if end_idx is None:
            return self.paragraphs[start_idx]["text"]

        end_idx = min(end_idx, len(self.paragraphs) - 1)
        return "\n".join([p["text"] for p in self.paragraphs[start_idx:end_idx + 1]])

    def extract_table_data(self, table_idx, row_start=0, row_end=None, col_start=0, col_end=None):
        """提取表格数据"""
        if not self.is_loaded or table_idx >= len(self.tables):
            return []

        table = self.tables[table_idx]
        data = table["data"]

        row_end = row_end if row_end is not None else len(data)
        row_end = min(row_end, len(data))

        result = []
        for row_idx in range(row_start, row_end):
            row = data[row_idx]
            col_end_effective = col_end if col_end is not None else len(row)
            col_end_effective = min(col_end_effective, len(row))

            row_result = []
            for col_idx in range(col_start, col_end_effective):
                row_result.append(row[col_idx]["text"])

            result.append(row_result)

        return result


class DocumentLoadWorker(QRunnable):
    """用于异步加载文档的工作线程"""

    class Signals(QObject):
        """工作线程信号"""
        finished = pyqtSignal(DocumentContent)
        error = pyqtSignal(str)

    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path
        self.signals = self.Signals()

    @pyqtSlot()
    def run(self):
        """线程执行函数"""
        try:
            doc_content = DocumentContent(self.file_path)
            if doc_content.is_loaded:
                self.signals.finished.emit(doc_content)
            else:
                self.signals.error.emit(doc_content.load_error)
        except Exception as e:
            self.signals.error.emit(f"加载文档时出错: {str(e)}")


class DocumentManager(QObject):
    """文档管理器，处理文档加载与内容管理"""

    documentLoaded = pyqtSignal(DocumentContent)
    loadError = pyqtSignal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.current_document = None
        self.thread_pool = QThreadPool()
        self.thread_pool.setMaxThreadCount(4)  # 限制最大线程数

    def load_document(self, file_path):
        """异步加载文档"""
        worker = DocumentLoadWorker(file_path)
        worker.signals.finished.connect(self._on_document_loaded)
        worker.signals.error.connect(self._on_load_error)
        self.thread_pool.start(worker)

    def _on_document_loaded(self, document):
        """文档加载完成回调"""
        self.current_document = document
        self.documentLoaded.emit(document)

    def _on_load_error(self, error):
        """文档加载错误回调"""
        self.loadError.emit(error)

    def get_current_document(self):
        """获取当前文档"""
        return self.current_document
